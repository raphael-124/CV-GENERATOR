from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.lib.colors import HexColor
from reportlab.platypus import Paragraph
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT


def _rgb_from_hex(hex_str):
    hex_str = (hex_str or '').lstrip('#')
    if len(hex_str) != 6:
        hex_str = '000000'
    return RGBColor(int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16))

def build_header(document, data):
    p = document.add_paragraph()
    run = p.add_run(data.get("name", "").upper())
    run.font.size = Pt(18)
    run.bold = True
    if data.get("job_title"):
        p = document.add_paragraph(data["job_title"]) 
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    contact = []
    for k in ["phone", "email", "location", "linkedin", "github", "website"]:
        v = data.get(k)
        if v:
            contact.append(v)
    if contact:
        document.add_paragraph(" | ".join(contact))

def add_heading(document, text, accent=None):
    h = document.add_paragraph()
    r = h.add_run(text.upper())
    r.bold = True
    if accent:
        r.font.color.rgb = _rgb_from_hex(accent)

def build_docx(data, output_path, template="sidebar", accent="#b87333"):
    document = Document()
    if template == "sidebar":
        add_heading(document, "Curriculum Vitae (CV)", accent)
        table = document.add_table(rows=1, cols=2)
        table.autofit = False
        left, right = table.rows[0].cells
        left.width = Inches(2.2)
        right.width = Inches(4.8)
        def shade(cell, color_hex):
            tcPr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), color_hex.replace('#',''))
            tcPr.append(shd)
        shade(left, accent)
        lp = left.add_paragraph()
        r = lp.add_run((data.get("name", "")).upper())
        r.bold = True
        r.font.size = Pt(16)
        if data.get("job_title"):
            left.add_paragraph(data["job_title"]).alignment = WD_ALIGN_PARAGRAPH.LEFT
        ph = left.add_paragraph("Contact Details")
        ph.runs[0].bold = True
        contact = []
        for k in ["phone","email","location","linkedin","github","website"]:
            v = data.get(k)
            if v:
                contact.append(v)
        for ctext in contact:
            left.add_paragraph(ctext)
        skills = data.get("skills", [])
        if skills:
            sh = left.add_paragraph("Core Skills")
            sh.runs[0].bold = True
            for s in skills:
                left.add_paragraph(s, style="List Bullet")
        languages = data.get("languages", [])
        if languages:
            lh = left.add_paragraph("Languages")
            lh.runs[0].bold = True
            left.add_paragraph(", ".join(languages))
        rp = right.add_paragraph("Professional Profile")
        rp.runs[0].bold = True
        rp.runs[0].font.color.rgb = _rgb_from_hex(accent)
        if data.get("summary"):
            sp = right.add_paragraph(data["summary"]) 
            sp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        experiences = data.get("experiences", [])
        if experiences:
            eh = right.add_paragraph("Career Summary")
            eh.runs[0].bold = True
            eh.runs[0].font.color.rgb = _rgb_from_hex(accent)
            for exp in experiences:
                head = []
                if exp.get("company"):
                    head.append(exp["company"]) 
                if exp.get("location"):
                    head.append(exp["location"]) 
                if exp.get("title"):
                    head.append(exp["title"]) 
                date = exp.get("dates") or ""
                if date:
                    right.add_paragraph(date)
                if head:
                    p = right.add_paragraph(" | ".join(head))
                    p.runs[0].bold = True
                for b in exp.get("bullets", []):
                    right.add_paragraph(b, style="List Bullet")
        education = data.get("education", [])
        if education:
            edh = right.add_paragraph("Education")
            edh.runs[0].bold = True
            edh.runs[0].font.color.rgb = _rgb_from_hex(accent)
            for ed in education:
                line = []
                if ed.get("degree"):
                    line.append(ed["degree"]) 
                if ed.get("institution"):
                    line.append(ed["institution"]) 
                p = right.add_paragraph(" | ".join(line))
                p.runs[0].bold = True
                meta = []
                if ed.get("location"):
                    meta.append(ed["location"]) 
                if ed.get("dates"):
                    meta.append(ed["dates"]) 
                if meta:
                    right.add_paragraph(" | ".join(meta))
        projects = data.get("projects", [])
        if projects:
            prh = right.add_paragraph("Projects")
            prh.runs[0].bold = True
            prh.runs[0].font.color.rgb = _rgb_from_hex(accent)
            for pr in projects:
                name = pr.get("name") or "Project"
                rp = right.add_paragraph(name)
                rp.runs[0].bold = True
                if pr.get("tech"):
                    right.add_paragraph(pr["tech"]) 
                for b in pr.get("bullets", []):
                    right.add_paragraph(b, style="List Bullet")
                if pr.get("link"):
                    right.add_paragraph(pr["link"]) 
        certifications = data.get("certifications", [])
        if certifications:
            ch = right.add_paragraph("Certifications")
            ch.runs[0].bold = True
            ch.runs[0].font.color.rgb = _rgb_from_hex(accent)
            for c in certifications:
                right.add_paragraph(c)
        extras = data.get("extras", [])
        if extras:
            xh = right.add_paragraph("Extras")
            xh.runs[0].bold = True
            xh.runs[0].font.color.rgb = _rgb_from_hex(accent)
            for e in extras:
                right.add_paragraph(e)
        if data.get("references"):
            rh = right.add_paragraph("References")
            rh.runs[0].bold = True
            rh.runs[0].font.color.rgb = _rgb_from_hex(accent)
            rp = document.add_paragraph(data["references"]) 
            rp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif template == "band":
        band = document.add_table(rows=1, cols=1)
        band.autofit = True
        cell = band.rows[0].cells[0]
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), accent.replace('#',''))
        tcPr.append(shd)
        p = cell.add_paragraph()
        r = p.add_run((data.get("name", "")).upper())
        r.bold = True
        r.font.size = Pt(18)
        r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        if data.get("job_title"):
            pr = cell.add_paragraph()
            rr = pr.add_run(data["job_title"]) 
            rr.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        contact = [v for k in ["phone","email","location","linkedin","github","website"] for v in ([data.get(k)] if data.get(k) else [])]
        if contact:
            pc = cell.add_paragraph(" | ".join(contact))
            pc.runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        add_heading(document, "Professional Profile", accent)
        if data.get("summary"):
            pp = document.add_paragraph(data["summary"]) 
            pp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        experiences = data.get("experiences", [])
        if experiences:
            add_heading(document, "Work Experience", accent)
            for exp in experiences:
                head = " | ".join([x for x in [exp.get("title"), exp.get("company")] if x])
                p = document.add_paragraph(head)
                p.runs[0].bold = True
                meta = " | ".join([x for x in [exp.get("location"), exp.get("dates")] if x])
                if meta:
                    document.add_paragraph(meta)
                for b in exp.get("bullets", []):
                    document.add_paragraph(b, style="List Bullet")
        education = data.get("education", [])
        if education:
            add_heading(document, "Education", accent)
            for ed in education:
                head = " | ".join([x for x in [ed.get("degree"), ed.get("institution")] if x])
                p = document.add_paragraph(head)
                p.runs[0].bold = True
                meta = " | ".join([x for x in [ed.get("location"), ed.get("dates")] if x])
                if meta:
                    document.add_paragraph(meta)
        projects = data.get("projects", [])
        if projects:
            add_heading(document, "Projects", accent)
            for pr in projects:
                rp = document.add_paragraph(pr.get("name") or "Project")
                rp.runs[0].bold = True
                if pr.get("tech"):
                    document.add_paragraph(pr["tech"]) 
                for b in pr.get("bullets", []):
                    document.add_paragraph(b, style="List Bullet")
                if pr.get("link"):
                    document.add_paragraph(pr["link"]) 
        certifications = data.get("certifications", [])
        if certifications:
            add_heading(document, "Certifications", accent)
            for c in certifications:
                document.add_paragraph(c)
        extras = data.get("extras", [])
        if extras:
            add_heading(document, "Extras", accent)
            for e in extras:
                document.add_paragraph(e)
        languages = data.get("languages", [])
        if languages:
            add_heading(document, "Languages", accent)
            document.add_paragraph(", ".join(languages))
        if data.get("references"):
            add_heading(document, "References", accent)
            rp = document.add_paragraph(data["references"]) 
            rp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        add_heading(document, "Curriculum Vitae (CV)", accent)
        namep = document.add_paragraph()
        r = namep.add_run((data.get("name", "")).upper())
        r.bold = True
        r.font.size = Pt(18)
        if data.get("job_title"):
            document.add_paragraph(data["job_title"]) 
        contacts = [data.get(k) for k in ["phone","email","location","linkedin","github","website"] if data.get(k)]
        if contacts:
            document.add_paragraph(" | ".join(contacts))
        if data.get("summary"):
            add_heading(document, "Summary", accent)
            sp = document.add_paragraph(data["summary"]) 
            sp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        skills = data.get("skills", [])
        if skills:
            add_heading(document, "Skills", accent)
            for s in skills:
                document.add_paragraph(s, style="List Bullet")
        experiences = data.get("experiences", [])
        if experiences:
            add_heading(document, "Work Experience", accent)
            for exp in experiences:
                head = " | ".join([x for x in [exp.get("title"), exp.get("company")] if x])
                p = document.add_paragraph(head)
                p.runs[0].bold = True
                meta = " | ".join([x for x in [exp.get("location"), exp.get("dates")] if x])
                if meta:
                    document.add_paragraph(meta)
                for b in exp.get("bullets", []):
                    document.add_paragraph(b, style="List Bullet")
        education = data.get("education", [])
        if education:
            add_heading(document, "Education", accent)
            for ed in education:
                head = " | ".join([x for x in [ed.get("degree"), ed.get("institution")] if x])
                p = document.add_paragraph(head)
                p.runs[0].bold = True
                meta = " | ".join([x for x in [ed.get("location"), ed.get("dates")] if x])
                if meta:
                    document.add_paragraph(meta)
        projects = data.get("projects", [])
        if projects:
            add_heading(document, "Projects", accent)
            for pr in projects:
                rp = document.add_paragraph(pr.get("name") or "Project")
                rp.runs[0].bold = True
                if pr.get("tech"):
                    document.add_paragraph(pr["tech"]) 
                for b in pr.get("bullets", []):
                    document.add_paragraph(b, style="List Bullet")
                if pr.get("link"):
                    document.add_paragraph(pr["link"]) 
        certifications = data.get("certifications", [])
        if certifications:
            add_heading(document, "Certifications", accent)
            for c in certifications:
                document.add_paragraph(c)
        extras = data.get("extras", [])
        if extras:
            add_heading(document, "Extras", accent)
            for e in extras:
                document.add_paragraph(e)
        languages = data.get("languages", [])
        if languages:
            add_heading(document, "Languages", accent)
            document.add_paragraph(", ".join(languages))
        if data.get("references"):
            add_heading(document, "References", accent)
            rp = document.add_paragraph(data["references"]) 
            rp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    document.save(output_path)

def build_pdf(data, output_path, template="sidebar", accent="#b87333"):
    c = canvas.Canvas(output_path, pagesize=LETTER)
    page_w, page_h = LETTER
    margin = 0.7*inch

    def wrap_draw(x, y_ref, text, font='Helvetica', size=11, color=HexColor('#000000'), width=page_w-2*margin, gap=14, justify=False):
        style = ParagraphStyle(
            name='Custom',
            fontName=font,
            fontSize=size,
            textColor=color,
            leading=gap,
            alignment=TA_JUSTIFY if justify else TA_LEFT,
        )
        p = Paragraph(text.replace('\n', '<br/>'), style)
        p_w, p_h = p.wrap(width, page_h)
        if y_ref[0] - p_h < margin:
            c.showPage()
            if template == 'sidebar':
                sidebar_w = 2.1*inch
                c.setFillColor(HexColor(accent))
                c.rect(0, 0, sidebar_w, page_h, stroke=0, fill=1)
            y_ref[0] = page_h - margin
            c.setFillColor(HexColor('#000000'))
        p.drawOn(c, x, y_ref[0] - p_h)
        y_ref[0] -= (p_h + 4)

    def draw_bullet_paragraph(x, y_ref, text, width, color=HexColor('#000000'), bullet_char='•', bullet_gap=12):
        style = ParagraphStyle(
            name='Bullet',
            fontName='Helvetica',
            fontSize=11,
            leading=14,
            leftIndent=bullet_gap,
            bulletIndent=0,
            textColor=color
        )
        p = Paragraph(text, style, bulletText=bullet_char)
        p_w, p_h = p.wrap(width, page_h)
        if y_ref[0] - p_h < margin:
            c.showPage()
            if template == 'sidebar':
                sidebar_w = 2.1*inch
                c.setFillColor(HexColor(accent))
                c.rect(0, 0, sidebar_w, page_h, stroke=0, fill=1)
            y_ref[0] = page_h - margin
            c.setFillColor(HexColor('#000000'))
        p.drawOn(c, x, y_ref[0] - p_h)
        y_ref[0] -= (p_h + 2)

    if template == 'sidebar':
        sidebar_w = 2.1*inch
        right_x = sidebar_w + margin
        right_w = page_w - right_x - margin
        c.setFillColor(HexColor(accent))
        c.rect(0, 0, sidebar_w, page_h, stroke=0, fill=1)
        c.setFillColor(HexColor('#000000'))
        ly = page_h - margin
        ry = page_h - margin
        def section_right(title):
            nonlocal ry
            wrap_draw(right_x, [ry], title.upper(), 'Helvetica-Bold', 12, width=right_w)
            ry -= 4
        c.setFillColor(HexColor('#ffffff'))
        c.setFont('Helvetica-Bold', 16)
        c.drawString(margin*0.6, ly, (data.get('name','')).upper())
        ly -= 18
        c.setFont('Helvetica', 12)
        if data.get('job_title'):
            c.drawString(margin*0.6, ly, data['job_title'])
            ly -= 16
        c.setFont('Helvetica-Bold', 12)
        c.drawString(margin*0.6, ly, 'Contact Details')
        ly -= 14
        c.setFont('Helvetica', 11)
        for k in ['phone','email','location','linkedin','github','website']:
            v = data.get(k)
            if v:
                wrap_draw(margin*0.6, [ly], v, 'Helvetica', 11, HexColor('#ffffff'), width=sidebar_w - margin)
                ly -= 2
        skills = data.get('skills', [])
        if skills:
            ly -= 6
            c.setFont('Helvetica-Bold', 12)
            c.drawString(margin*0.6, ly, 'Core Skills')
            ly -= 14
            for s in skills:
                draw_bullet_paragraph(margin*0.6, [ly], s, width=sidebar_w - margin - 12, color=HexColor('#ffffff'))
                ly -= 2
        languages = data.get('languages', [])
        if languages:
            ly -= 6
            c.setFont('Helvetica-Bold', 12)
            c.drawString(margin*0.6, ly, 'Languages')
            ly -= 14
            wrap_draw(margin*0.6, [ly], ', '.join(languages), 'Helvetica', 11, HexColor('#ffffff'), width=sidebar_w - margin)
        c.setFillColor(HexColor('#000000'))
        wrap_draw(right_x, [ry], 'Curriculum Vitae (CV)', 'Helvetica-Bold', 14, width=right_w)
        ry -= 8
        wrap_draw(right_x, [ry], (data.get('name','')).upper(), 'Helvetica-Bold', 16, width=right_w)
        ry -= 10
        if data.get('job_title'):
            wrap_draw(right_x, [ry], data['job_title'], 'Helvetica', 12, width=right_w)
            ry -= 6
        contacts = [data.get(k) for k in ['phone','email','location','linkedin','github','website'] if data.get(k)]
        if contacts:
            wrap_draw(right_x, [ry], ' | '.join(contacts), 'Helvetica', 11, width=right_w)
            ry -= 4
        if data.get('summary'):
            section_right('Professional Profile')
            wrap_draw(right_x, [ry], data['summary'], 'Helvetica', 11, width=right_w, justify=True)
            ry -= 6
        experiences = data.get('experiences', [])
        if experiences:
            section_right('Career Summary')
            for exp in experiences:
                head = ' | '.join([x for x in [exp.get('company'), exp.get('location'), exp.get('title')] if x])
                date = exp.get('dates') or ''
                if date:
                    wrap_draw(right_x, [ry], date, 'Helvetica', 11, width=right_w)
                if head:
                    wrap_draw(right_x, [ry], head, 'Helvetica-Bold', 11, width=right_w)
                for b in exp.get('bullets', []):
                    draw_bullet_paragraph(right_x, [ry], b, width=right_w)
                ry -= 6
        education = data.get('education', [])
        if education:
            section_right('Education')
            for ed in education:
                head = ' | '.join([x for x in [ed.get('degree'), ed.get('institution')] if x])
                meta = ' | '.join([x for x in [ed.get('location'), ed.get('dates')] if x])
                if head:
                    wrap_draw(right_x, [ry], head, 'Helvetica-Bold', 11, width=right_w)
                if meta:
                    wrap_draw(right_x, [ry], meta, 'Helvetica', 11, width=right_w)
                ry -= 6
        projects = data.get('projects', [])
        if projects:
            section_right('Projects')
            for pr in projects:
                wrap_draw(right_x, [ry], pr.get('name') or 'Project', 'Helvetica-Bold', 11, width=right_w)
                if pr.get('tech'):
                    wrap_draw(right_x, [ry], pr['tech'], 'Helvetica', 11, width=right_w)
                for b in pr.get('bullets', []):
                    draw_bullet_paragraph(right_x, [ry], b, width=right_w)
                if pr.get('link'):
                    wrap_draw(right_x, [ry], pr['link'], 'Helvetica', 11, width=right_w)
                ry -= 6
        certifications = data.get('certifications', [])
        if certifications:
            section_right('Certifications')
            for ctext in certifications:
                wrap_draw(right_x, [ry], ctext, 'Helvetica', 11, width=right_w)
            ry -= 6
        extras = data.get('extras', [])
        if extras:
            section_right('Extras')
            for e in extras:
                wrap_draw(right_x, [ry], e, 'Helvetica', 11, width=right_w)
            ry -= 6
        languages = data.get('languages', [])
        if languages:
            section_right('Languages')
            wrap_draw(right_x, [ry], ', '.join(languages), 'Helvetica', 11, width=right_w)
            ry -= 6
        if data.get('references'):
            section_right('References')
            wrap_draw(right_x, [ry], data['references'], 'Helvetica', 11, width=right_w, justify=True)
    elif template == 'band':
        band_h = 0.9*inch
        c.setFillColor(HexColor(accent))
        c.rect(0, page_h - band_h, page_w, band_h, stroke=0, fill=1)
        y = page_h - margin
        c.setFillColor(HexColor('#ffffff'))
        c.setFont('Helvetica-Bold', 18)
        c.drawString(margin, y-4, (data.get('name','')).upper())
        c.setFont('Helvetica', 12)
        if data.get('job_title'):
            c.drawString(margin, y-22, data['job_title'])
        contacts = [data.get(k) for k in ['phone','email','location','linkedin','github','website'] if data.get(k)]
        if contacts:
            c.setFont('Helvetica', 10)
            wrap_draw(margin, [y-38], ' | '.join(contacts), 'Helvetica', 10, HexColor('#ffffff'), width=page_w-2*margin)
        c.setFillColor(HexColor('#000000'))
        y = page_h - band_h - margin
        def section(title):
            nonlocal y
            wrap_draw(margin, [y], title.upper(), 'Helvetica-Bold', 12)
            y -= 4
        if data.get('summary'):
            section('Professional Profile')
            wrap_draw(margin, [y], data['summary'], 'Helvetica', 11, justify=True)
            y -= 6
        exps = data.get('experiences', [])
        if exps:
            section('Work Experience')
            for exp in exps:
                head = ' | '.join([x for x in [exp.get('title'), exp.get('company')] if x])
                wrap_draw(margin, [y], head, 'Helvetica-Bold', 11)
                meta = ' | '.join([x for x in [exp.get('location'), exp.get('dates')] if x])
                if meta:
                    wrap_draw(margin, [y], meta, 'Helvetica', 11)
                for b in exp.get('bullets', []):
                    draw_bullet_paragraph(margin, [y], b, width=page_w-2*margin)
                y -= 6
        edu = data.get('education', [])
        if edu:
            section('Education')
            for ed in edu:
                head = ' | '.join([x for x in [ed.get('degree'), ed.get('institution')] if x])
                wrap_draw(margin, [y], head, 'Helvetica-Bold', 11)
                meta = ' | '.join([x for x in [ed.get('location'), ed.get('dates')] if x])
                if meta:
                    wrap_draw(margin, [y], meta, 'Helvetica', 11)
                y -= 6
        prj = data.get('projects', [])
        if prj:
            section('Projects')
            for pr in prj:
                wrap_draw(margin, [y], pr.get('name') or 'Project', 'Helvetica-Bold', 11)
                if pr.get('tech'):
                    wrap_draw(margin, [y], pr['tech'], 'Helvetica', 11)
                for b in pr.get('bullets', []):
                    draw_bullet_paragraph(margin, [y], b, width=page_w-2*margin)
                if pr.get('link'):
                    wrap_draw(margin, [y], pr['link'], 'Helvetica', 11)
                y -= 6
        certs = data.get('certifications', [])
        if certs:
            section('Certifications')
            for ctext in certs:
                wrap_draw(margin, [y], ctext, 'Helvetica', 11)
            y -= 6
        extras = data.get('extras', [])
        if extras:
            section('Extras')
            for e in extras:
                wrap_draw(margin, [y], e, 'Helvetica', 11)
            y -= 6
        languages = data.get('languages', [])
        if languages:
            section('Languages')
            wrap_draw(margin, [y], ', '.join(languages), 'Helvetica', 11)
            y -= 6
        if data.get('references'):
            section('References')
            wrap_draw(margin, [y], data['references'], 'Helvetica', 11, justify=True)
    else:
        y = page_h - margin
        def section(title):
            nonlocal y
            wrap_draw(margin, [y], title.upper(), 'Helvetica-Bold', 12)
            y -= 4
        wrap_draw(margin, [y], 'Curriculum Vitae (CV)', 'Helvetica-Bold', 14)
        wrap_draw(margin, [y], (data.get('name','')).upper(), 'Helvetica-Bold', 16)
        if data.get('job_title'):
            wrap_draw(margin, [y], data['job_title'], 'Helvetica', 12)
        contacts = [data.get(k) for k in ['phone','email','location','linkedin','github','website'] if data.get(k)]
        if contacts:
            wrap_draw(margin, [y], ' | '.join(contacts), 'Helvetica', 11)
        if data.get('summary'):
            section('Summary')
            wrap_draw(margin, [y], data['summary'], 'Helvetica', 11, justify=True)
        skills = data.get('skills', [])
        if skills:
            section('Skills')
            for s in skills:
                draw_bullet_paragraph(margin, [y], s, width=page_w-2*margin)
        exps = data.get('experiences', [])
        if exps:
            section('Work Experience')
            for exp in exps:
                head = ' | '.join([x for x in [exp.get('title'), exp.get('company')] if x])
                wrap_draw(margin, [y], head, 'Helvetica-Bold', 11)
                meta = ' | '.join([x for x in [exp.get('location'), exp.get('dates')] if x])
                if meta:
                    wrap_draw(margin, [y], meta, 'Helvetica', 11)
                for b in exp.get('bullets', []):
                    draw_bullet_paragraph(margin, [y], b, width=page_w-2*margin)
        edu = data.get('education', [])
        if edu:
            section('Education')
            for ed in edu:
                head = ' | '.join([x for x in [ed.get('degree'), ed.get('institution')] if x])
                wrap_draw(margin, [y], head, 'Helvetica-Bold', 11)
                meta = ' | '.join([x for x in [ed.get('location'), ed.get('dates')] if x])
                if meta:
                    wrap_draw(margin, [y], meta, 'Helvetica', 11)
        prj = data.get('projects', [])
        if prj:
            section('Projects')
            for pr in prj:
                wrap_draw(margin, [y], pr.get('name') or 'Project', 'Helvetica-Bold', 11)
                if pr.get('tech'):
                    wrap_draw(margin, [y], pr['tech'], 'Helvetica', 11)
                for b in pr.get('bullets', []):
                    draw_bullet_paragraph(margin, [y], b, width=page_w-2*margin)
                if pr.get('link'):
                    wrap_draw(margin, [y], pr['link'], 'Helvetica', 11)
        certs = data.get('certifications', [])
        if certs:
            section('Certifications')
            for ctext in certs:
                wrap_draw(margin, [y], ctext, 'Helvetica', 11)
        extras = data.get('extras', [])
        if extras:
            section('Extras')
            for e in extras:
                wrap_draw(margin, [y], e, 'Helvetica', 11)
        languages = data.get('languages', [])
        if languages:
            section('Languages')
            wrap_draw(margin, [y], ', '.join(languages), 'Helvetica', 11)
        if data.get('references'):
            section('References')
            wrap_draw(margin, [y], data['references'], 'Helvetica', 11, justify=True)
    c.setFillColor(HexColor('#000000'))
    c.setFont('Helvetica', 9)
    c.drawString(page_w/2 - 14, margin/2, 'Page 1')
    c.showPage()
    c.save()